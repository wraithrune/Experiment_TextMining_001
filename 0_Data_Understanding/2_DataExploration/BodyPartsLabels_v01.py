# -*- coding: utf-8 -*-
"""
Created on Sun Oct  1 21:58:07 2017

@author: Wei Jing
"""

# --- Construction Labels Creator ---
# --- Objective: Identify if an accident case is construction or non-construction 

# --- Extract Text from PDFs
# --- Import libraries
import PyPDF2
import docx

import pandas
import string
import nltk
from nltk import word_tokenize, FreqDist
from nltk.corpus import stopwords

# --- Import global variables
gRawConstructionTerms = []
gProcessedConstructionTerms = []
gConstructionTermAppearCount = []

# --- Function for Text Extraction from PDFs
def fTxtExtractorPDF(vPDFFilePath, vExtractPageRange, vExtractStartPage):
    # Local Variables
    vPDF_TotalPages = 0 # Contains no. of pages in PDF file
    vContent = "" # Contains string of text to extract in PDF pages
    
    # PDF Code Segment
    vPDFFileObj = open(vPDFFilePath,'rb') # Open PDF in vPDFFilePath in read binary mode
    vPDFReader = PyPDF2.PdfFileReader(vPDFFileObj) # Pass vPDFFileObj to vPDFReader for PDF representation
    #vPDF_TotalPages = vPDFReader.numPages # Store the no. of pages in PDF
    #print("Total Pages in PDF is: ", vPDF_TotalPages)
    
    if vExtractPageRange > 0: # Begin Loop if the page range to extract is not zero
        for i in range(vExtractPageRange):
            # Extract the pages in the page range one by one
            vContent += vPDFReader.getPage((vExtractStartPage + i - 1)).extractText() + "\n"
            #vContent
            #print("Page ", (vExtractStartPage + i - 1), " contains the following text: ", vContent)
    
    # Word Docx Code Segment
    vDoc = docx.Document()
    vDoc.add_paragraph(vContent)
    vDoc.save('PDFtoWord.docx')
    
    return vContent

# Call Function fTxtExtractorPDF
# Actual pages for Construction terms pg. 84-86
fTxtExtractorPDF("oiics_manual_2010_a.pdf", 4, 83)

# --- Function for Full Text Extraction from Word Document
def fTxtExtractorWordDocx(vWordDocxFilePath):
    # Local Variables
    vFullText = []
    vDoc = docx.Document(vWordDocxFilePath)
    
    for vParagraph in vDoc.paragraphs: # Loops over all Paragraph objects
        vFullText.append(vParagraph.text) # Appends Paragraph text to list
    
    print(str(vFullText))
    
    return str(vFullText)

# --- Function for Corpus Preprocessing (Tokenize, Remove punctuations and stopwords, Apply wtem)
def fLoadCorpus4Process(vCorpus):
    #print("Corpus",vCorpus)
    # Convert the Free Text into Tokens
    vTokens = nltk.word_tokenize(str(vCorpus)) #word_tokenize(vCorpus)
    #print("Tokens", vTokens)
    # Remove Punctuations
    vTokens_nop = [ t for t in vTokens if t not in string.punctuation]
    
    # Convert All Lower Case
    vTokens_lower=[ t.lower() for t in vTokens_nop]
    
    # Create Stopwords
    vStopWords = stopwords.words('english')
	
    #select english stopwords
    vStopWords = set(stopwords.words("english"))
    
    # Remove All Stopwords From Text
    vTokens_nostop=[ t for t in vTokens_lower if t not in vStopWords]
    #print(vTokens_nostop)
    
    # Apply Porter Stemmer
    vPorter = nltk.PorterStemmer()
    vTokens_porter=[vPorter.stem(t) for t in vTokens_nostop] 
    
    print(vTokens_porter)
    
    return vTokens_porter

# --- Function for Creating Bag of Words from Excel from 1 Column
def fBagOfWords(vExcelFilePath, vColumnName):
    
    vColumnLen = 0
    vColumnStr = ""
    vText = []
    
    # Load Excel Corpus
    vExcelCorpus = pandas.read_excel(vExcelFilePath)
    
    vColumnLen = len(vExcelCorpus[vColumnName])
    for i in range (vColumnLen):
        vText.append(vExcelCorpus[vColumnName][i])
        
    #print("vText", vText[2])
        
    # Convert the Free Text into Tokens
    vCTokens = word_tokenize(str(vText))
    
    # Remove Punctuations
    vTokens_nop = [ t for t in vCTokens if t not in string.punctuation]
    
    # Convert All Lower Case
    vTokens_lower=[ t.lower() for t in vTokens_nop]
    
    # Create Stopwords
    vStopWords = stopwords.words('english')
	
	 #select english stopwords
    vStopWords = set(stopwords.words("english"))
    
    # Remove All Stopwords From Text
    vTokens_nostop=[ t for t in vTokens_lower if t not in vStopWords]
    
    # Apply Porter Stemmer
    vPorter = nltk.PorterStemmer()
    vTokens_porter=[vPorter.stem(t) for t in vTokens_nostop]
    
    return
    
# --- Function for checking if an accident case is construction related
def fConstructionCase(vExcelFilePath, vColumnName, vConstructionDict):
    
    vCorpusRowCount = []
    vCounter = 0
    vText = []
    
    # Load Excel Corpus
    vExcelCorpus = pandas.read_excel(vExcelFilePath)
    
    vColumnLen = len(vExcelCorpus[vColumnName])
    for i in range (vColumnLen):
        vText.append(vExcelCorpus[vColumnName][i])
        
    vDictionaryLen = len(vConstructionDict)
    for j in range (vColumnLen):
        for k in range (vDictionaryLen):
            #print("str(vConstructionDict[k])", str(vConstructionDict[k]))
            #print("str(vText[j]",str(vText[j]))
            if str(vConstructionDict[k]) in str(vText[j]):
                vCounter += 1
            if k == vDictionaryLen-1:
                vCorpusRowCount.append(vCounter)
                vCounter = 0
                print("This row contains count of:",vCorpusRowCount[j])

    return vCorpusRowCount

# --- Function for Output to Excel (hardcode for now)
def fOutput(vCounterArray):
    
    vConstructTF = ""
    
    # Load Excel Corpus
    vExcelCorpus = pandas.read_excel("MsiaAccidentCases.xlsx")
    
    f = open('OSHA_TEST_ConstructionLabels.csv', 'w', newline="\n") # open a csv file for writing
    # --- Write CSV headers
    f.write("Cause" + ',' + "Title Case" + ',' + "Construction Dictionary Counts" + ',' + "Construction T/F" + ',' + "Summary Case" + "\n")
    
    vCountlen = len(vCounterArray)
    
    for i in range (vCountlen):
        if vCounterArray[i] != 0:
            vConstructTF = "TRUE"
        else:
            vConstructTF = "FALSE"
            
        f.write(str(vExcelCorpus["Cause"][i]) + ',' + str(vExcelCorpus["Title Case"][i]) + ',' + str(vCounterArray[i]) + ',' + vConstructTF + ',' + str(vExcelCorpus["Summary Case"][i]) + "\n")

#gRawConstructionTerms = fTxtExtractorWordDocx("PDFtoWord.docx")
#gProcessedConstructionTerms = fLoadCorpus4Process(gRawConstructionTerms)
#fBagOfWords("MsiaAccidentCases.xlsx", "Summary Case")
#print("gRawConstructionTerms", gProcessedConstructionTerms)
#gConstructionTermAppearCount = fConstructionCase("MsiaAccidentCases.xlsx", "Summary Case", gProcessedConstructionTerms)
#gConstructionTermAppearCount = fConstructionCase("osha_test.xlsx", "Summary Case", gProcessedConstructionTerms)
#fOutput(gConstructionTermAppearCount)
# --- Heuristic

    